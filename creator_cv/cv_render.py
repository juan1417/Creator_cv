from __future__ import annotations

import asyncio
import html
import io
import json
import os
import platform
import re
from pathlib import Path
from typing import Any

import markdown as md_lib
from docx import Document
from fpdf import FPDF
try:
    from fpdf.enums import XPos, YPos
except Exception:  # pragma: no cover - compatibilidad con paquetes fpdf antiguos
    XPos = None
    YPos = None


def markdown_to_preview_html(md: str) -> str:
    """HTML seguro para vista previa en navegador (Markdown → HTML)."""
    return md_lib.markdown(
        md,
        extensions=["extra", "nl2br", "sane_lists"],
    )


def _e(value: Any) -> str:
    return html.escape("" if value is None else str(value))


def _lines(s: str) -> list[str]:
    return [ln.strip() for ln in s.splitlines() if ln.strip()] if s else []


def _normalize_url(link: str) -> str:
    """Si la URL no tiene esquema, le antepone https://."""
    s = (link or "").strip()
    if not s:
        return ""
    if "://" in s:
        return s
    if s.startswith("//"):
        return "https:" + s
    return "https://" + s


def _fmt_date_short(value: str) -> str:
    """
    Formatea fechas 'YYYY-MM' o 'YYYY-MM-DD' como 'Mon YYYY' o 'YYYY'.
    Devuelve el string original si no encaja.
    """
    if not value:
        return ""
    v = str(value).strip()
    parts = v.split("-")
    months = {
        "01": "Jan", "02": "Feb", "03": "Mar", "04": "Apr",
        "05": "May", "06": "Jun", "07": "Jul", "08": "Aug",
        "09": "Sep", "10": "Oct", "11": "Nov", "12": "Dec",
    }
    if len(parts) >= 2 and parts[1] in months:
        return f"{months[parts[1]]} {parts[0]}"
    return v


def _normalize_link_list(links: Any) -> list[str]:
    """Lista de enlaces o un solo string con URLs separadas por comas."""
    if links is None:
        return []
    if isinstance(links, list):
        result: list[str] = []
        for x in links:
            s = str(x).strip()
            if not s:
                continue
            for part in s.split(","):
                p = part.strip()
                if p:
                    result.append(p)
        return result
    if isinstance(links, str):
        s = links.strip()
        if not s:
            return []
        return [p.strip() for p in s.split(",") if p.strip()]
    return []


def _clean_url_for_compare(url: str) -> str:
    """Limpia una URL para comparar dos enlaces (ignora scheme y www)."""
    u = url.strip().lower()
    for prefix in ("https://", "http://", "www."):
        if u.startswith(prefix):
            u = u[len(prefix):]
    return u.rstrip("/")


def _contact_line(
    meta: dict[str, Any],
    data: dict[str, Any] | None = None,
) -> list[tuple[str, str]]:
    """Pares (etiqueta, valor) para la cabecera; solo entradas con texto."""
    raw = meta.get("contacto")
    if not isinstance(raw, dict):
        return []
    order = (
        ("Teléfono", "telefono"),
        ("Email", "email"),
        ("LinkedIn", "linkedin"),
        ("Ubicación", "ubicacion"),
    )
    out: list[tuple[str, str]] = []
    for label, key in order:
        val = raw.get(key)
        if val is None:
            continue
        s = str(val).strip()
        if s:
            out.append((label, s))

    # Agregar portafolio desde recursos_actuales.links
    if data:
        rec = data.get("recursos_actuales") or {}
        links = _normalize_link_list(rec.get("links"))
        if links:
            linkedin_raw = (raw.get("linkedin") or "").strip()
            linkedin_clean = _clean_url_for_compare(linkedin_raw)
            for link in links:
                link_stripped = link.strip()
                if not link_stripped:
                    continue
                if _clean_url_for_compare(link_stripped) == linkedin_clean:
                    continue
                out.append(("Portafolio", link_stripped))

    return out


def _meta_contact_nonempty(meta: dict[str, Any]) -> bool:
    return bool(_contact_line(meta))


def _merge_tecnicas_lists(
    explicit: Any,
    from_keywords: Any,
) -> list[str]:
    """
    Lista única para mostrar/exportar: habilidades.tecnicas primero, luego
    perfil_profesional.palabras_clave que no estén ya (mismo uso práctico que stack).
    """
    out: list[str] = []
    seen: set[str] = set()

    def _add_many(raw: Any) -> None:
        if raw is None:
            return
        items = raw if isinstance(raw, list) else [raw]
        for x in items:
            s = str(x).strip()
            if not s:
                continue
            key = s.casefold()
            if key not in seen:
                seen.add(key)
                out.append(s)

    _add_many(explicit)
    _add_many(from_keywords)
    return out


def context_has_preview_content(data: dict[str, Any]) -> bool:
    """True si la vista previa estructurada tendría bloques de CV visibles."""
    meta = data.get("meta") or {}
    if str(meta.get("nombre_completo") or "").strip():
        return True
    if str(meta.get("titulo_profesional") or "").strip():
        return True
    if str(meta.get("objetivo_cv") or "").strip():
        return True
    if _meta_contact_nonempty(meta):
        return True
    perfil = data.get("perfil_profesional") or {}
    if (perfil.get("resumen") or "").strip():
        return True
    if perfil.get("palabras_clave"):
        return True
    if data.get("experiencia"):
        return True
    if data.get("educacion"):
        return True
    hab = data.get("habilidades") or {}
    if hab.get("tecnicas") or hab.get("blandas") or hab.get("idiomas"):
        return True
    if data.get("proyectos"):
        return True
    if data.get("certificaciones"):
        return True
    if data.get("fortalezas"):
        return True
    return False


def _exp_date_range(item: dict[str, Any]) -> str:
    fin = item.get("fecha_fin") or ("Present" if item.get("actual") else "")
    parts = [_fmt_date_short(x) for x in (item.get("fecha_inicio") or "", fin) if x]
    return " – ".join(parts)


def context_to_structured_preview_html(
    data: dict[str, Any],
    *,
    fallback_title: str = "",
) -> str:
    """
    HTML semántico para vista previa (estilo Harvard: una columna, lineal,
    texto en negro, jerarquía por peso y tamaño).
    Todo el texto de usuario pasa por escape HTML.
    """
    parts: list[str] = []
    meta = data.get("meta") or {}
    nombre = (meta.get("nombre_completo") or "").strip()
    if not nombre and fallback_title:
        nombre = fallback_title.strip()

    parts.append('<article class="cv-ref">')

    parts.append('<header class="cv-ref__header">')
    if nombre:
        parts.append(f'<h1 class="cv-ref__name">{_e(nombre)}</h1>')
    contact = _contact_line(meta, data)
    if contact:
        parts.append('<p class="cv-ref__contact-line">')
        for i, (label, val) in enumerate(contact):
            if i > 0:
                parts.append('<span class="cv-ref__sep" aria-hidden="true">|</span> ')
            if label == "Portafolio":
                url = _normalize_url(val)
                parts.append(
                    f'<span class="cv-ref__contact-val">'
                    f'<a href="{_e(url)}" target="_blank" rel="noopener noreferrer" '
                    f'class="cv-ref__proj-link-anchor">{_e(val)}</a>'
                    f'</span>'
                )
            else:
                parts.append(f'<span class="cv-ref__contact-val">{_e(val)}</span>')
        parts.append("</p>")
    parts.append("</header>")

    perfil = data.get("perfil_profesional") or {}
    resumen = (perfil.get("resumen") or "").strip()
    if resumen:
        parts.append('<section class="cv-ref__section">')
        parts.append('<h2 class="cv-ref__section-title">Perfil Profesional</h2>')
        parts.append(f'<p class="cv-ref__para">{_e(resumen)}</p>')
        parts.append("</section>")

    exp = data.get("experiencia") or []
    if isinstance(exp, list) and exp:
        parts.append('<section class="cv-ref__section">')
        parts.append('<h2 class="cv-ref__section-title">Experiencia</h2>')
        for item in exp:
            if not isinstance(item, dict):
                parts.append(f'<p class="cv-ref__entry">{_e(item)}</p>')
                continue
            cargo = (item.get("cargo") or "").strip()
            org = (item.get("empresa") or "").strip()
            loc = (item.get("ubicacion") or "").strip()
            drange = _exp_date_range(item)
            parts.append('<div class="cv-ref__entry">')
            parts.append('<div class="cv-ref__entry-head">')
            parts.append('<div class="cv-ref__entry-main">')
            if cargo:
                parts.append(f'<span class="cv-ref__entry-role">{_e(cargo)}</span>')
            if org:
                parts.append(f'<span class="cv-ref__entry-org">, {_e(org)}</span>')
            parts.append("</div>")
            parts.append('<div class="cv-ref__entry-aside">')
            if drange:
                parts.append(f'<span class="cv-ref__entry-date">{_e(drange)}</span>')
            if loc:
                parts.append(f'<span class="cv-ref__entry-loc">{_e(loc)}</span>')
            parts.append("</div></div>")
            for key in ("responsabilidades", "logros"):
                val = item.get(key)
                if not val:
                    continue
                rows = val if isinstance(val, list) else _lines(str(val))
                if rows:
                    parts.append('<ul class="cv-ref__bullets">')
                    for row in rows:
                        parts.append(f"<li>{_e(row)}</li>")
                    parts.append("</ul>")
            parts.append("</div>")
        parts.append("</section>")

    hab = data.get("habilidades") or {}
    perfil_kw = (perfil.get("palabras_clave") or [])
    tech = _merge_tecnicas_lists(hab.get("tecnicas"), perfil_kw)
    soft = hab.get("blandas") or []
    langs = hab.get("idiomas") or []
    if tech or soft or langs:
        parts.append('<section class="cv-ref__section">')
        parts.append('<h2 class="cv-ref__section-title">Habilidades</h2>')
        if tech:
            parts.append(
                '<p class="cv-ref__skills-row">'
                f'<span class="cv-ref__skills-label">Técnicas:</span> '
                f'<span class="cv-ref__skills-val">{_e(", ".join(str(x) for x in tech))}</span>'
                "</p>"
            )
        if soft:
            parts.append(
                '<p class="cv-ref__skills-row">'
                f'<span class="cv-ref__skills-label">Habilidades Blandas:</span> '
                f'<span class="cv-ref__skills-val">{_e(", ".join(str(x) for x in soft))}</span>'
                "</p>"
            )
        if langs:
            parts.append(
                '<p class="cv-ref__skills-row">'
                f'<span class="cv-ref__skills-label">Idiomas:</span> '
                f'<span class="cv-ref__skills-val">{_e(", ".join(str(x) for x in langs))}</span>'
                "</p>"
            )
        parts.append("</section>")

    edu = data.get("educacion") or []
    certs = data.get("certificaciones") or []
    has_edu = isinstance(edu, list) and bool(edu)
    has_certs = isinstance(certs, list) and bool(certs)
    if has_edu or has_certs:
        parts.append('<section class="cv-ref__section">')
        parts.append('<h2 class="cv-ref__section-title">Educación y Certificaciones</h2>')
        if has_edu:
            for item in edu:
                if not isinstance(item, dict):
                    parts.append(f'<p class="cv-ref__entry">{_e(item)}</p>')
                    continue
                titulo = (item.get("titulo") or "").strip()
                inst = (item.get("institucion") or "").strip()
                loc = (item.get("ubicacion") or "").strip()
                drange_parts: list[str] = []
                if item.get("fecha_inicio") or item.get("fecha_fin"):
                    drange_parts.append(
                        " – ".join(
                            x
                            for x in (
                                item.get("fecha_inicio") or "",
                                item.get("fecha_fin") or "Present",
                            )
                            if x
                        )
                    )
                if item.get("estado"):
                    drange_parts.append(str(item["estado"]))
                drange = " · ".join(drange_parts)
                parts.append('<div class="cv-ref__entry">')
                parts.append('<div class="cv-ref__entry-head">')
                parts.append('<div class="cv-ref__entry-main">')
                if inst:
                    parts.append(f'<span class="cv-ref__entry-org">{_e(inst)}</span>')
                if titulo:
                    if inst:
                        parts.append(" — ")
                    parts.append(f'<span class="cv-ref__entry-role">{_e(titulo)}</span>')
                parts.append("</div>")
                parts.append('<div class="cv-ref__entry-aside">')
                if drange:
                    parts.append(f'<span class="cv-ref__entry-date">{_e(drange)}</span>')
                if loc:
                    parts.append(f'<span class="cv-ref__entry-loc">{_e(loc)}</span>')
                parts.append("</div></div>")
                parts.append("</div>")
        if has_certs:
            for c in certs:
                if not isinstance(c, dict):
                    parts.append(f'<p class="cv-ref__entry">{_e(c)}</p>')
                    continue
                name = (c.get("nombre") or c.get("titulo") or "").strip()
                desc = (c.get("descripcion") or c.get("detalle") or "").strip()
                if not name and not desc:
                    continue
                parts.append('<div class="cv-ref__entry">')
                if name:
                    parts.append(f'<span class="cv-ref__entry-role">{_e(name)}</span>')
                if desc:
                    parts.append(
                        f'<span class="cv-ref__entry-org"> — {_e(desc)}</span>'
                    )
                parts.append("</div>")
        parts.append("</section>")

    proy = data.get("proyectos") or []
    if isinstance(proy, list) and proy:
        parts.append('<section class="cv-ref__section">')
        parts.append('<h2 class="cv-ref__section-title">Proyectos</h2>')
        for item in proy:
            if not isinstance(item, dict):
                parts.append(f'<p class="cv-ref__entry">{_e(item)}</p>')
                continue
            name = (item.get("nombre") or "").strip()
            desc = (item.get("descripcion") or "").strip()
            tec = item.get("tecnologias") or []
            link = (item.get("enlace") or "").strip()
            parts.append('<div class="cv-ref__entry">')
            if name:
                parts.append(
                    f'<div class="cv-ref__proj-name">{_e(name)}</div>'
                )
            has_meta = desc or (isinstance(tec, list) and tec) or link
            if has_meta:
                parts.append('<div class="cv-ref__proj-meta">')
                if desc:
                    parts.append(f'<span class="cv-ref__proj-desc">{_e(desc)}</span>')
                if isinstance(tec, list) and tec:
                    parts.append(
                        '<span class="cv-ref__proj-tech">'
                        f'<span class="cv-ref__proj-tech-label">Tecnologías:</span> '
                        f'<span class="cv-ref__proj-tech-val">{_e(", ".join(str(x) for x in tec))}</span>'
                        "</span>"
                    )
                if link:
                    href = _normalize_url(link)
                    parts.append(
                        '<span class="cv-ref__proj-link">'
                        f'<span class="cv-ref__proj-link-label">Enlace:</span> '
                        f'<a class="cv-ref__proj-link-anchor" href="{_e(href)}" target="_blank" rel="noopener noreferrer">{_e(link)}</a>'
                        "</span>"
                    )
                parts.append("</div>")
            parts.append("</div>")
        parts.append("</section>")

    strengths = data.get("fortalezas") or []
    if isinstance(strengths, list) and strengths:
        rows: list[str] = []
        for s in strengths:
            if not isinstance(s, dict):
                v = (str(s) or "").strip()
                if v:
                    rows.append(v)
                continue
            st = (s.get("titulo") or s.get("nombre") or "").strip()
            sd = (s.get("descripcion") or "").strip()
            if st and sd:
                rows.append(f"{st}: {sd}")
            elif st:
                rows.append(st)
            elif sd:
                rows.append(sd)
        if rows:
            parts.append('<section class="cv-ref__section">')
            parts.append('<h2 class="cv-ref__section-title">Adicional</h2>')
            parts.append('<ul class="cv-ref__bullets">')
            for row in rows:
                parts.append(f"<li>{_e(row)}</li>")
            parts.append("</ul>")
            parts.append("</section>")

    parts.append("</article>")
    return "".join(parts)


def context_to_preview_document_html(
    data: dict[str, Any],
    *,
    css_text: str,
    fallback_title: str = "",
) -> str:
    """Documento HTML completo para exportación fiel del preview."""
    preview_html = context_to_structured_preview_html(data, fallback_title=fallback_title)
    doc_title = _e(fallback_title or "CV")
    return (
        "<!DOCTYPE html>"
        '<html lang="es"><head><meta charset="utf-8" />'
        '<meta name="viewport" content="width=device-width, initial-scale=1" />'
        f"<title>{doc_title}</title>"
        f"<style>{css_text}</style>"
        "<style>"
        "body{margin:0;background:#fff;color:#111827;}"
        ".preview-doc-stage{padding:0;background:transparent;border-radius:0;}"
        ".cv-paper{max-width:none;min-height:auto;box-shadow:none;border-radius:0;}"
        ".cv-paper-inner.cv-ref-doc.cv-ref-doc--pdf{padding:10mm 12mm;font-size:10pt;}"
        ".cv-ref-doc--pdf .cv-ref{break-inside:auto !important;page-break-inside:auto !important;}"
        ".cv-ref-doc--pdf .cv-ref__header{break-after:auto !important;page-break-after:auto !important;}"
        "/* Layout PDF con columnas fijas y paginacion estable */"
        ".cv-ref-doc--pdf .cv-ref__grid{display:table !important;width:100% !important;table-layout:fixed !important;}"
        ".cv-ref-doc--pdf .cv-ref__aside,.cv-ref-doc--pdf .cv-ref__main{display:table-cell !important;vertical-align:top !important;overflow:visible !important;}"
        ".cv-ref-doc--pdf .cv-ref__aside{width:32% !important;padding-right:1.2rem !important;}"
        ".cv-ref-doc--pdf .cv-ref__main{width:68% !important;}"
        ".cv-ref-doc--pdf .cv-ref__section,.cv-ref-doc--pdf .cv-ref__timeline-item,.cv-ref-doc--pdf .cv-ref__edu-item{break-inside:auto !important;page-break-inside:auto !important;}"
        "@page{size:A4;margin:0;}"
        "</style></head><body>"
        '<div class="preview-doc-stage"><div class="cv-paper">'
        '<div class="cv-paper-inner cv-ref-doc cv-ref-doc--pdf">'
        f"{preview_html}"
        "</div></div></div></body></html>"
    )


def _css_from_app(root_path: str) -> str:
    css_path = Path(root_path) / "static" / "css" / "app.css"
    try:
        return css_path.read_text(encoding="utf-8")
    except OSError as e:
        raise RuntimeError(f"No se pudo leer CSS del preview: {css_path}") from e


def _render_html_to_pdf_with_weasyprint(html_doc: str) -> bytes:
    try:
        from weasyprint import HTML
    except Exception as e:  # pragma: no cover - dependencia opcional en runtime
        raise RuntimeError(
            "WeasyPrint no está disponible. Instala con: uv add weasyprint"
        ) from e
    try:
        return HTML(string=html_doc).write_pdf()
    except Exception as e:
        raise RuntimeError(f"WeasyPrint no pudo renderizar el PDF: {e}") from e


async def _render_html_to_pdf_with_playwright(html_doc: str) -> bytes:
    try:
        from playwright.async_api import async_playwright
    except Exception as e:  # pragma: no cover - dependencia opcional en runtime
        raise RuntimeError(
            "Playwright no está disponible. Instala con: uv add playwright "
            "y luego: uv run playwright install chromium"
        ) from e

    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch()
            try:
                page = await browser.new_page()
                await page.set_content(html_doc, wait_until="networkidle")
                return await page.pdf(
                    format="A4",
                    print_background=True,
                    margin={"top": "0mm", "right": "0mm", "bottom": "0mm", "left": "0mm"},
                )
            finally:
                await browser.close()
    except Exception as e:
        raise RuntimeError(
            f"Playwright falló al renderizar: {e}. "
            "Ejecuta: uv run playwright install chromium"
        ) from e


def context_to_pdf_bytes_from_preview(
    data: dict[str, Any],
    *,
    root_path: str,
    fallback_title: str = "",
) -> bytes:
    """
    Renderiza PDF desde el mismo HTML/CSS del preview (fiel al estilo/íconos).
    """
    css_text = _css_from_app(root_path)
    html_doc = context_to_preview_document_html(
        data,
        css_text=css_text,
        fallback_title=fallback_title,
    )
    # 1) WeasyPrint suele paginar columnas mejor para CV largos.
    try:
        return _render_html_to_pdf_with_weasyprint(html_doc)
    except RuntimeError:
        # 2) Fallback a Chromium/Playwright.
        return asyncio.run(_render_html_to_pdf_with_playwright(html_doc))


def _fmt_val(v: Any) -> str:
    if v is None:
        return ""
    if isinstance(v, bool):
        return "sí" if v else "no"
    if isinstance(v, (list, dict)):
        return json.dumps(v, ensure_ascii=False)
    return str(v)


def json_to_markdown(data: dict[str, Any]) -> str:
    """Solo formatea el JSON ya guardado; no inventa contenido."""
    lines: list[str] = []

    meta = data.get("meta") or {}
    if any(meta.values()):
        lines.append("## Meta")
        for key in (
            "nombre_completo",
            "titulo_profesional",
            "idioma_cv",
            "objetivo_cv",
            "tipo_cv",
            "nivel_seniority",
        ):
            val = meta.get(key)
            if val:
                lines.append(f"- **{key}**: {val}")
        contact = meta.get("contacto")
        if isinstance(contact, dict):
            contact_parts: list[str] = []
            for ck in ("telefono", "email", "linkedin", "ubicacion"):
                cv = contact.get(ck)
                if cv and str(cv).strip():
                    contact_parts.append(str(cv).strip())
            # Agregar portafolio desde recursos_actuales.links (dedup LinkedIn)
            rec = data.get("recursos_actuales") or {}
            linkedin_raw = (contact.get("linkedin") or "").strip()
            for link in _normalize_link_list(rec.get("links")):
                link_stripped = link.strip()
                if link_stripped and _clean_url_for_compare(link_stripped) != _clean_url_for_compare(linkedin_raw):
                    contact_parts.append(link_stripped)
            if contact_parts:
                lines.append("- " + " | ".join(contact_parts))
        lines.append("")

    perfil = data.get("perfil_profesional") or {}
    resumen = (perfil.get("resumen") or "").strip()
    if resumen:
        lines.append("## Perfil Profesional")
        lines.append(resumen)
        lines.append("")

    exp = data.get("experiencia") or []
    if exp:
        lines.append("## Experiencia")
        for item in exp:
            if not isinstance(item, dict):
                lines.append(f"- {_fmt_val(item)}")
                continue
            title = item.get("cargo") or ""
            org = item.get("empresa") or ""
            head = ", ".join(x for x in (title, org) if x)
            if head:
                lines.append(f"### {head}")
            dates = " – ".join(
                _fmt_date_short(x)
                for x in (
                    item.get("fecha_inicio") or "",
                    item.get("fecha_fin") or ("Present" if item.get("actual") else ""),
                )
                if x
            )
            if dates:
                lines.append(dates)
            for label, key in (
                ("Ubicación", "ubicacion"),
                ("Responsabilidades", "responsabilidades"),
                ("Logros", "logros"),
            ):
                val = item.get(key)
                if not val:
                    continue
                if isinstance(val, list):
                    lines.append(f"**{label}:**")
                    for row in val:
                        lines.append(f"- {row}")
                else:
                    lines.append(f"**{label}:** {val}")
            lines.append("")

    hab = data.get("habilidades") or {}
    if hab:
        perfil_kw = (data.get("perfil_profesional") or {}).get("palabras_clave") or []
        tech = _merge_tecnicas_lists(hab.get("tecnicas"), perfil_kw)
        soft = hab.get("blandas") or []
        langs = hab.get("idiomas") or []
        if tech or soft or langs:
            lines.append("## Habilidades")
            if tech:
                lines.append("**Técnicas:** " + ", ".join(str(x) for x in tech))
            if soft:
                lines.append("**Habilidades Blandas:** " + ", ".join(str(x) for x in soft))
            if langs:
                lines.append("**Idiomas:** " + ", ".join(str(x) for x in langs))
            lines.append("")

    edu = data.get("educacion") or []
    certs = data.get("certificaciones") or []
    has_edu = isinstance(edu, list) and bool(edu)
    has_certs = isinstance(certs, list) and bool(certs)
    if has_edu or has_certs:
        lines.append("## Educación y Certificaciones")
        if has_edu:
            for item in edu:
                if not isinstance(item, dict):
                    lines.append(f"- {_fmt_val(item)}")
                    continue
                t = item.get("titulo") or ""
                inst = item.get("institucion") or ""
                head = " — ".join(x for x in (inst, t) if x)
                line = f"- {head}" if head else f"- {_fmt_val(item)}"
                extra = []
                if item.get("fecha_inicio") or item.get("fecha_fin"):
                    extra.append(
                        " – ".join(
                            x
                            for x in (
                                _fmt_date_short(item.get("fecha_inicio") or ""),
                                _fmt_date_short(item.get("fecha_fin") or "Present"),
                            )
                            if x
                        )
                    )
                if item.get("estado"):
                    extra.append(str(item["estado"]))
                if extra:
                    line += " (" + "; ".join(extra) + ")"
                lines.append(line)
        if has_certs:
            for c in certs:
                if not isinstance(c, dict):
                    lines.append(f"- {_fmt_val(c)}")
                    continue
                name = (c.get("nombre") or c.get("titulo") or "").strip()
                desc = (c.get("descripcion") or c.get("detalle") or "").strip()
                if name and desc:
                    lines.append(f"- **{name}**: {desc}")
                elif name:
                    lines.append(f"- **{name}**")
                elif desc:
                    lines.append(f"- {desc}")
        lines.append("")

    proy = data.get("proyectos") or []
    if proy:
        lines.append("## Proyectos")
        for item in proy:
            if not isinstance(item, dict):
                lines.append(f"- {_fmt_val(item)}")
                continue
            name = (item.get("nombre") or "Proyecto").strip()
            lines.append(f"### {name}")
            if item.get("descripcion"):
                lines.append(str(item["descripcion"]))
            if item.get("tecnologias"):
                tec = item["tecnologias"]
                if isinstance(tec, list):
                    lines.append("**Tecnologías:** " + ", ".join(str(x) for x in tec))
            if item.get("enlace"):
                link = str(item["enlace"]).strip()
                lines.append(f"**Enlace:** <{_normalize_url(link)}>")
            lines.append("")

    strengths = data.get("fortalezas") or []
    if isinstance(strengths, list) and strengths:
        lines.append("## Adicional")
        for s in strengths:
            if not isinstance(s, dict):
                lines.append(f"- {_fmt_val(s)}")
                continue
            st = (s.get("titulo") or s.get("nombre") or "").strip()
            sd = (s.get("descripcion") or "").strip()
            if st and sd:
                lines.append(f"- **{st}**: {sd}")
            elif st:
                lines.append(f"- **{st}**")
            elif sd:
                lines.append(f"- {sd}")
        lines.append("")

    rec = data.get("recursos_actuales") or {}
    # Mostrar solo si tiene cv_existente o texto_cv (links se muestran en contacto)
    has_meta = any(rec.get(k) for k in ("cv_existente", "texto_cv"))
    if has_meta:
        lines.append("## Recursos actuales")
        if rec.get("cv_existente"):
            lines.append(f"- **cv_existente:** {_fmt_val(rec.get('cv_existente'))}")
        if texto := rec.get("texto_cv"):
            lines.append("**Texto CV:**")
            lines.append(str(texto))
        lines.append("")

    rest = data.get("restricciones") or {}
    if any(v not in (None, "") for v in rest.values()):
        lines.append("## Restricciones")
        for key in (
            "extension_maxima_paginas",
            "formato_solicitado",
            "otro",
        ):
            val = rest.get(key)
            if val:
                lines.append(f"- **{key}:** {val}")
        lines.append("")

    dudas = data.get("dudas_pendientes") or []
    if dudas:
        lines.append("## Dudas pendientes")
        for d in dudas:
            lines.append(f"- {d}")
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def markdown_to_docx_bytes(md: str) -> bytes:
    doc = Document()
    for line in md.split("\n"):
        line = line.rstrip()
        if not line:
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _unicode_font_path() -> Path | None:
    """TTF con soporte Unicode (Windows / Linux / macOS)."""
    system = platform.system()
    candidates: list[Path] = []
    if system == "Windows":
        windir = os.environ.get("WINDIR", r"C:\Windows")
        candidates = [
            Path(windir) / "Fonts" / "arial.ttf",
            Path(windir) / "Fonts" / "calibri.ttf",
            Path(windir) / "Fonts" / "segoeui.ttf",
        ]
    elif system == "Darwin":
        candidates = [
            Path("/Library/Fonts/Arial.ttf"),
            Path("/System/Library/Fonts/Supplemental/Arial.ttf"),
        ]
    else:
        candidates = [
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            Path("/usr/share/fonts/TTF/DejaVuSans.ttf"),
            Path("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"),
        ]
    for p in candidates:
        if p.is_file():
            return p
    return None


def _strip_inline_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    return text


def _fold_latin_fallback(text: str) -> str:
    """Último recurso si no hay TTF: aproximar a Latin-1."""
    return (
        text.replace("\u2014", "-")
        .replace("\u2013", "-")
        .replace("\u202f", " ")
        .replace("\u00a0", " ")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u2018", "'")
        .replace("\u2019", "'")
    )


def context_to_pdf_bytes(
    data: dict[str, Any],
    *,
    fallback_title: str = "",
) -> bytes:
    """
    PDF con plantilla Harvard de una sola columna, en negro, mismo orden
    de secciones que la vista previa HTML/MD. Sin dependencias externas
    más allá de fpdf2.
    """
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=14)
    pdf.set_margins(18, 16, 18)
    pdf.add_page()

    font_path = _unicode_font_path()
    use_custom_font = bool(font_path)
    if use_custom_font:
        pdf.add_font("CvBody", "", str(font_path))
        font = "CvBody"
    else:
        font = "Helvetica"

    page_left = pdf.l_margin
    page_right = pdf.w - pdf.r_margin
    page_width = pdf.epw
    black = (0, 0, 0)

    # Tamaños (mm) — equilibrados para A4
    name_sz = 22
    contact_sz = 10
    section_sz = 10
    role_sz = 11.5
    title_sz = 10.5
    body_sz = 10
    small_sz = 9.5
    lh = 5.0
    bullet_lh = 4.6

    def txt(v: Any) -> str:
        s = "" if v is None else str(v)
        return _fold_latin_fallback(s) if not font_path else s

    def set_font(size: float, *, bold: bool = False) -> None:
        style = ""
        if bold and not use_custom_font:
            style = "B"
        pdf.set_font(font, style, size=size)

    def write_block(
        x: float,
        y: float,
        w: float,
        text: str,
        size: float,
        *,
        align: str = "L",
        line_height: float = lh,
    ) -> float:
        text = txt(text).strip()
        if not text:
            return y
        set_font(size)
        pdf.set_xy(x, y)
        if XPos is not None and YPos is not None:
            pdf.multi_cell(
                w,
                line_height,
                text,
                align=align,
                new_x=XPos.LEFT,
                new_y=YPos.NEXT,
            )
        else:
            pdf.multi_cell(w, line_height, text, align=align)
        return pdf.get_y()

    def section_title(x: float, y: float, w: float, title: str) -> float:
        pdf.set_text_color(*black)
        set_font(section_sz, bold=True)
        pdf.set_xy(x, y)
        t = txt(title).upper()
        if XPos is not None and YPos is not None:
            pdf.multi_cell(w, lh, t, new_x=XPos.LEFT, new_y=YPos.NEXT)
        else:
            pdf.multi_cell(w, lh, t)
        y_after = pdf.get_y()
        pdf.set_draw_color(*black)
        pdf.set_line_width(0.25)
        pdf.line(x, y_after + 0.6, x + w, y_after + 0.6)
        return y_after + 2.6

    def entry_head(
        x: float,
        y: float,
        w: float,
        left_text: str,
        right_text: str,
        *,
        left_bold: bool = False,
    ) -> float:
        """Cabecera de entrada: título a la izquierda, fecha a la derecha."""
        pdf.set_text_color(*black)
        if not right_text:
            return write_block(
                x, y, w, left_text,
                title_sz if left_bold else body_sz,
                line_height=lh,
            )
        right_w = max(35.0, w * 0.32)
        left_w = w - right_w - 4
        right_render = txt(right_text).strip()
        if not right_render:
            return write_block(
                x, y, w, left_text,
                title_sz if left_bold else body_sz,
                line_height=lh,
            )
        # 1) Escribir primero la columna izquierda, partiendo de Y original.
        left_y = write_block(
            x, y, left_w, left_text,
            title_sz if left_bold else body_sz,
            line_height=lh,
        )
        # 2) Volver a Y original y escribir la columna derecha.
        pdf.set_xy(x + left_w + 4, y)
        set_font(small_sz)
        if XPos is not None and YPos is not None:
            pdf.multi_cell(
                right_w,
                lh,
                right_render,
                align="R",
                new_x=XPos.LEFT,
                new_y=YPos.NEXT,
            )
        else:
            pdf.multi_cell(right_w, lh, right_render, align="R")
        right_y = pdf.get_y()
        return max(left_y, right_y)

    def write_bullets(
        x: float,
        y: float,
        w: float,
        rows: list[str],
        size: float = body_sz,
    ) -> float:
        cur_y = y
        for row in rows:
            s = txt(row).strip()
            if not s:
                continue
            set_font(size)
            pdf.set_text_color(*black)
            pdf.set_xy(x, cur_y)
            line = f"• {s}"
            if XPos is not None and YPos is not None:
                pdf.multi_cell(
                    w, bullet_lh, line,
                    new_x=XPos.LEFT, new_y=YPos.NEXT,
                )
            else:
                pdf.multi_cell(w, bullet_lh, line)
            cur_y = pdf.get_y()
        return cur_y

    # ─────────────────────── HEADER ───────────────────────
    meta = data.get("meta") or {}
    nombre = (meta.get("nombre_completo") or "").strip() or fallback_title.strip()
    contact = _contact_line(meta, data)

    pdf.set_text_color(*black)
    y = pdf.t_margin
    if nombre:
        set_font(name_sz, bold=True)
        pdf.set_xy(page_left, y)
        nm = txt(nombre)
        if XPos is not None and YPos is not None:
            pdf.multi_cell(page_width, 9, nm, align="C",
                           new_x=XPos.LEFT, new_y=YPos.NEXT)
        else:
            pdf.multi_cell(page_width, 9, nm, align="C")
        y = pdf.get_y() + 1.5

    if contact:
        set_font(contact_sz)
        pdf.set_xy(page_left, y)
        parts = []
        for label, val in contact:
            parts.append(f"{label}: {val}" if label and label != "Portafolio" else str(val))
        line = " | ".join(parts)
        line = txt(line)
        if XPos is not None and YPos is not None:
            pdf.multi_cell(page_width, lh, line, align="C",
                           new_x=XPos.LEFT, new_y=YPos.NEXT)
        else:
            pdf.multi_cell(page_width, lh, line, align="C")
        y = pdf.get_y() + 2

    # Línea bajo la cabecera
    pdf.set_draw_color(*black)
    pdf.set_line_width(0.4)
    pdf.line(page_left, y, page_right, y)
    y += 4

    # ─────────────────────── PROFILE ───────────────────────
    perfil = data.get("perfil_profesional") or {}
    resumen = (perfil.get("resumen") or "").strip()
    if resumen:
        y = section_title(page_left, y, page_width, "Perfil Profesional")
        y = write_block(page_left, y, page_width, resumen, body_sz, line_height=lh)
        y += 2

    # ─────────────────────── EXPERIENCE ───────────────────────
    exp = data.get("experiencia") or []
    if isinstance(exp, list) and exp:
        y = section_title(page_left, y, page_width, "Experiencia")
        for item in exp:
            if not isinstance(item, dict):
                y = write_block(page_left, y, page_width, str(item), body_sz, line_height=lh)
                continue
            cargo = (item.get("cargo") or "").strip()
            org = (item.get("empresa") or "").strip()
            loc = (item.get("ubicacion") or "").strip()
            drange = _exp_date_range(item)
            title = ", ".join(x for x in (cargo, org) if x)
            if title:
                y = entry_head(page_left, y, page_width, title, drange, left_bold=True) + 0.5
            elif drange:
                y = write_block(page_left, y, page_width, drange, small_sz, line_height=lh)
            for key in ("responsabilidades", "logros"):
                val = item.get(key)
                if not val:
                    continue
                rows = val if isinstance(val, list) else _lines(str(val))
                if rows:
                    y = write_bullets(page_left, y + 0.3, page_width, rows, body_sz)
            if loc:
                y = write_block(page_left, y, page_width, loc, small_sz, line_height=lh)
            y += 2

    # ─────────────────────── SKILLS ───────────────────────
    hab = data.get("habilidades") or {}
    perfil_kw = (perfil.get("palabras_clave") or [])
    tech = _merge_tecnicas_lists(hab.get("tecnicas"), perfil_kw)
    soft = hab.get("blandas") or []
    langs = hab.get("idiomas") or []
    if tech or soft or langs:
        y = section_title(page_left, y, page_width, "Habilidades")
        if tech:
            line = "Técnicas: " + ", ".join(str(x) for x in tech)
            y = write_block(page_left, y, page_width, line, body_sz, line_height=lh) + 0.3
        if soft:
            line = "Habilidades Blandas: " + ", ".join(str(x) for x in soft)
            y = write_block(page_left, y, page_width, line, body_sz, line_height=lh) + 0.3
        if langs:
            line = "Idiomas: " + ", ".join(str(x) for x in langs)
            y = write_block(page_left, y, page_width, line, body_sz, line_height=lh)
        y += 2

    # ─────────────────────── EDUCATION & CERTIFICATIONS ───────────────────────
    edu = data.get("educacion") or []
    certs = data.get("certificaciones") or []
    has_edu = isinstance(edu, list) and bool(edu)
    has_certs = isinstance(certs, list) and bool(certs)
    if has_edu or has_certs:
        y = section_title(page_left, y, page_width, "Educación y Certificaciones")
        if has_edu:
            for item in edu:
                if not isinstance(item, dict):
                    y = write_block(page_left, y, page_width, str(item), body_sz, line_height=lh)
                    continue
                titulo = (item.get("titulo") or "").strip()
                inst = (item.get("institucion") or "").strip()
                loc = (item.get("ubicacion") or "").strip()
                left = " — ".join(x for x in (inst, titulo) if x)
                # rango de fechas + estado
                dr_parts: list[str] = []
                if item.get("fecha_inicio") or item.get("fecha_fin"):
                    fin_raw = item.get("fecha_fin") or "Present"
                    dr_parts.append(
                        " – ".join(
                            _fmt_date_short(x)
                            for x in (item.get("fecha_inicio") or "", fin_raw)
                            if x
                        )
                    )
                if item.get("estado"):
                    dr_parts.append(str(item["estado"]))
                right = " · ".join(dr_parts)
                if left:
                    y = entry_head(page_left, y, page_width, left, right, left_bold=False) + 0.5
                elif right:
                    y = write_block(page_left, y, page_width, right, small_sz, line_height=lh)
                if loc:
                    y = write_block(page_left, y, page_width, loc, small_sz, line_height=lh)
                y += 1.5
        if has_certs:
            for c in certs:
                if not isinstance(c, dict):
                    y = write_block(page_left, y, page_width, str(c), body_sz, line_height=lh)
                    continue
                name = (c.get("nombre") or c.get("titulo") or "").strip()
                desc = (c.get("descripcion") or c.get("detalle") or "").strip()
                left = f"{name} — {desc}" if name and desc else (name or desc)
                if left:
                    y = write_block(page_left, y, page_width, left, body_sz, line_height=lh) + 1
        y += 1

    # ─────────────────────── PROJECTS ───────────────────────
    proy = data.get("proyectos") or []
    if isinstance(proy, list) and proy:
        y = section_title(page_left, y, page_width, "Proyectos")
        for item in proy:
            if not isinstance(item, dict):
                y = write_block(page_left, y, page_width, str(item), body_sz, line_height=lh)
                continue
            name = (item.get("nombre") or "").strip()
            desc = (item.get("descripcion") or "").strip()
            tec = item.get("tecnologias") or []
            link = (item.get("enlace") or "").strip()
            if name:
                set_font(role_sz, bold=True)
                pdf.set_text_color(*black)
                y = write_block(page_left, y, page_width, name, role_sz, line_height=lh)
            if desc:
                y = write_block(page_left, y, page_width, desc, body_sz, line_height=lh)
            if isinstance(tec, list) and tec:
                y = write_block(
                    page_left, y, page_width,
                    f"Tecnologías: {', '.join(str(t) for t in tec)}",
                    small_sz, line_height=lh,
                )
            if link:
                href = _normalize_url(link)
                y = write_block(
                    page_left, y, page_width,
                    f"Enlace: {href}",
                    small_sz, line_height=lh,
                )
            y += 2

    # ─────────────────────── ADDITIONAL (fortalezas) ───────────────────────
    strengths = data.get("fortalezas") or []
    if isinstance(strengths, list) and strengths:
        rows: list[str] = []
        for s in strengths:
            if not isinstance(s, dict):
                v = (str(s) or "").strip()
                if v:
                    rows.append(v)
                continue
            st = (s.get("titulo") or s.get("nombre") or "").strip()
            sd = (s.get("descripcion") or "").strip()
            if st and sd:
                rows.append(f"{st}: {sd}")
            elif st:
                rows.append(st)
            elif sd:
                rows.append(sd)
        if rows:
            y = section_title(page_left, y, page_width, "Adicional")
            y = write_bullets(page_left, y, page_width, rows, body_sz)

    return bytes(pdf.output())


def markdown_to_pdf_bytes(md: str) -> bytes:
    """
    Compatibilidad temporal: genera un PDF básico desde texto Markdown.
    """
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=14)
    pdf.set_margins(14, 14, 14)
    pdf.add_page()
    font_path = _unicode_font_path()
    if font_path:
        pdf.add_font("CvBody", "", str(font_path))
        pdf.set_font("CvBody", size=10)
    else:
        pdf.set_font("Helvetica", size=10)
        md = _fold_latin_fallback(md)
    for raw in md.split("\n"):
        line = _strip_inline_markdown(raw.rstrip())
        if not line:
            pdf.ln(2)
            continue
        pdf.set_x(pdf.l_margin)
        if XPos is not None and YPos is not None:
            pdf.multi_cell(pdf.epw, 5, line, new_x=XPos.LEFT, new_y=YPos.NEXT)
        else:
            pdf.multi_cell(pdf.epw, 5, line)
    return bytes(pdf.output())
